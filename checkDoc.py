import os
import docx
import re

doc = docx.Document('INT-MCS-18018_DTSG_U2.docx')
versionRegex = re.compile(r'v\d.(\d)+')
# mo = versionRegex.search('(v1.0)\n')
# print('Version: '+mo.group())
# for j in range(0,58):
	# print('Predict: '+str(j)+' -- ' +doc.paragraphs[j].text)
logFile = open(r'.\log.txt','w')

# print('Len: '+str(len(doc.paragraphs)))
# p_xml = [paragraph._element.xml for paragraph in doc.paragraphs]
# for para in p_xml:
	# print(para+'\n')
	# try:
		# logFile.write(para+'\n')
	# except UnicodeEncodeError:
		# logFile.write('KhanhLog123: Error!\n')
# logFile.close()

for j in range(0,len(doc.paragraphs)-1):
	try:
		print(doc.paragraphs[j].text+'\n')
	except UnicodeEncodeError:
		print('Error!\n')
	
# index = 0
# checkParagraph = versionRegex.search(doc.paragraphs[index].text)
# while versionRegex.search(doc.paragraphs[index].text) == None:
	# index += 1
	
# version = versionRegex.search(doc.paragraphs[index].text).group()
# print('Version: '+version+' - index = '+str(index))